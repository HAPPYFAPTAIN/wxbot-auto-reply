# -*- coding: utf-8 -*-
"""
wxbot_files.py — 微信文件消息读取与解析。

- 存储目录：C:\\Users\\Administrator\\xwechat_files\\<wxid>_xxxx\\msg\\file\\YYYY-MM\\<文件名>
  （对方发来的文件微信会自动下载到这里；自己发的也在这）
- find_file(filename)：在最近两个月的目录里找同名文件，取最新
- parse_file(path, max_chars)：按扩展名解析 md/txt/docx/pdf/xlsx/xls/csv/json/py 等，返回纯文本（截断）
"""
import os, re, glob, time

TEXT_EXTS = {".md", ".txt", ".py", ".json", ".csv", ".log", ".yaml", ".yml", ".xml", ".html", ".js", ".ts", ".ini", ".cfg", ".sh", ".bat"}
DOCX_EXTS = {".docx"}
PDF_EXTS = {".pdf"}
XLSX_EXTS = {".xlsx", ".xlsm"}
XLS_EXTS = {".xls"}

_STORAGE_ROOT = r"C:\Users\Administrator\xwechat_files"


def _storage_dirs():
    """所有账号的 msg/file 目录，按最近修改排序（新的在前）。"""
    dirs = []
    try:
        for acc in os.listdir(_STORAGE_ROOT):
            d = os.path.join(_STORAGE_ROOT, acc, "msg", "file")
            if os.path.isdir(d):
                dirs.append(d)
    except Exception:
        pass
    return dirs


def find_file(filename, max_age_days=62):
    """在最近两个月的 YYYY-MM 子目录里找同名文件（大小写不敏感），返回最新路径或 None。"""
    filename = (filename or "").strip()
    if not filename:
        return None
    cands = []
    now = time.time()
    for d in _storage_dirs():
        # 只扫最近三个月目录，避免全量遍历
        subs = []
        try:
            subs = [os.path.join(d, s) for s in os.listdir(d) if re.match(r"^\d{4}-\d{2}$", s)]
        except Exception:
            continue
        subs.sort(reverse=True)
        for sd in subs[:3]:
            p = os.path.join(sd, filename)
            if os.path.exists(p):
                cands.append(p)
            else:
                # 微信可能对重名文件加 (1) 后缀
                base, ext = os.path.splitext(filename)
                for hit in glob.glob(os.path.join(sd, base + "*" + ext)):
                    if os.path.basename(hit).lower().startswith(base.lower()[:8]):
                        cands.append(hit)
    cands = [c for c in cands if now - os.path.getmtime(c) < max_age_days * 86400]
    if not cands:
        return None
    cands.sort(key=lambda p: os.path.getmtime(p), reverse=True)
    return cands[0]


def _trunc(text, max_chars):
    text = (text or "").strip()
    if len(text) > max_chars:
        return text[:max_chars] + f"\n…（内容过长已截断，共约{len(text)}字）"
    return text


def parse_file(path, max_chars=1500):
    """按扩展名解析文件为纯文本。失败返回错误说明字符串。"""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext in TEXT_EXTS:
            for enc in ("utf-8", "gbk", "utf-16", "latin-1"):
                try:
                    with open(path, "r", encoding=enc) as f:
                        return _trunc(f.read(), max_chars)
                except (UnicodeDecodeError, UnicodeError):
                    continue
            return "（文本编码无法识别）"
        if ext in DOCX_EXTS:
            import docx
            doc = docx.Document(path)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            # 表格也带上
            for tbl in doc.tables[:5]:
                for row in tbl.rows[:20]:
                    parts.append(" | ".join(c.text.strip() for c in row.cells))
            return _trunc("\n".join(parts), max_chars)
        if ext == ".doc":
            return "（.doc 旧格式暂不支持解析，请对方转 docx）"
        if ext in PDF_EXTS:
            try:
                import pdfplumber
                parts = []
                with pdfplumber.open(path) as pdf:
                    for page in pdf.pages[:8]:
                        t = page.extract_text() or ""
                        if t.strip():
                            parts.append(t)
                if parts:
                    return _trunc("\n".join(parts), max_chars)
            except Exception:
                pass
            import pypdf
            reader = pypdf.PdfReader(path)
            parts = [(p.extract_text() or "") for p in reader.pages[:8]]
            text = "\n".join(t for t in parts if t.strip())
            return _trunc(text, max_chars) if text.strip() else "（PDF 无文本层，可能是扫描件）"
        if ext in XLSX_EXTS:
            import openpyxl
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            parts = []
            for ws in wb.worksheets[:3]:
                parts.append(f"【表:{ws.title}】")
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i >= 30:
                        parts.append("…")
                        break
                    vals = ["" if v is None else str(v) for v in row]
                    if any(v.strip() for v in vals):
                        parts.append(" | ".join(vals))
            wb.close()
            return _trunc("\n".join(parts), max_chars)
        if ext in XLS_EXTS:
            import xlrd
            wb = xlrd.open_workbook(path)
            parts = []
            for ws in wb.sheets()[:3]:
                parts.append(f"【表:{ws.name}】")
                for i in range(min(ws.nrows, 30)):
                    vals = [str(v) for v in ws.row_values(i)]
                    if any(v.strip() for v in vals):
                        parts.append(" | ".join(vals))
            return _trunc("\n".join(parts), max_chars)
        return f"（暂不支持的文件类型 {ext or '未知'}）"
    except Exception as e:
        return f"（解析出错: {type(e).__name__} {e}）"


def filename_from_bubble(text):
    """从文件气泡文本提取文件名。格式：'文件\\n[进度: x%\\n]文件名.ext\\n大小\\n微信电脑版'"""
    lines = [l.strip() for l in (text or "").splitlines() if l.strip()]
    for l in lines[1:]:
        if l.startswith("进度") or l in ("上传中", "微信电脑版", "已下载"):
            continue
        if re.search(r"\.[A-Za-z0-9]{1,8}$", l):
            return l
    return lines[1] if len(lines) > 1 else ""
